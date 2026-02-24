"""Document parsing for Knowledge Base.

Supports parsing of:
- Markdown (.md)
- Text (.txt)
- PDF (.pdf) - via pypdf
- DOCX (.docx) - via python-docx
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Optional
from dataclasses import dataclass


@dataclass
class ParsedDocument:
    """Result of document parsing."""

    title: str
    content: str
    metadata: dict[str, Any]
    links: list[str]


class DocumentParser:
    """Parse various document formats."""

    def __init__(self):
        self._pdf_parser = None
        self._docx_parser = None

    async def parse_file(self, file_path: str) -> ParsedDocument:
        """Parse a document file.

        Args:
            file_path: Path to the file

        Returns:
            ParsedDocument with title, content, metadata, links
        """
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext == ".md":
            return await self._parse_markdown(file_path)
        elif ext == ".txt":
            return await self._parse_text(file_path)
        elif ext == ".pdf":
            return await self._parse_pdf(file_path)
        elif ext == ".docx":
            return await self._parse_docx(file_path)
        else:
            return await self._parse_text(file_path)

    async def parse_content(
        self,
        content: str,
        format: str = "markdown",
    ) -> ParsedDocument:
        """Parse document from content string.

        Args:
            content: Document content
            format: Format type (markdown, text, html)

        Returns:
            ParsedDocument
        """
        if format == "markdown":
            return self._parse_markdown_content(content)
        elif format == "html":
            return self._parse_html_content(content)
        else:
            return self._parse_text_content(content)

    async def _parse_markdown(self, file_path: str) -> ParsedDocument:
        """Parse markdown file."""
        path = Path(file_path)
        content = path.read_text(encoding="utf-8", errors="ignore")
        return self._parse_markdown_content(content, path.stem)

    def _parse_markdown_content(
        self, content: str, default_title: str = "Untitled"
    ) -> ParsedDocument:
        """Parse markdown content."""
        lines = content.split("\n")
        title = default_title
        links = []

        # Extract title from first heading
        for line in lines:
            if line.startswith("# "):
                title = line[2:].strip()
                break

        # Extract links
        link_pattern = re.compile(r"\[([^\]]+)\]\(([^\)]+)\)")
        for match in link_pattern.finditer(content):
            links.append(match.group(2))

        # Clean markdown
        cleaned = self._clean_markdown(content)

        return ParsedDocument(
            title=title,
            content=cleaned,
            metadata={"format": "markdown", "original_length": len(content)},
            links=links,
        )

    def _clean_markdown(self, text: str) -> str:
        """Remove markdown syntax, keep text."""
        # Remove headers
        text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
        # Remove bold/italic
        text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
        text = re.sub(r"\*([^*]+)\*", r"\1", text)
        text = re.sub(r"__([^_]+)__", r"\1", text)
        text = re.sub(r"_([^_]+)_", r"\1", text)
        # Remove code blocks
        text = re.sub(r"```[\s\S]*?```", "", text)
        text = re.sub(r"`([^`]+)`", r"\1", text)
        # Remove links but keep text
        text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
        # Remove images
        text = re.sub(r"!\[([^\]]*)\]\([^\)]+\)", r"\1", text)
        # Clean extra whitespace
        text = re.sub(r"\n{3,}", "\n\n", text)

        return text.strip()

    async def _parse_text(self, file_path: str) -> ParsedDocument:
        """Parse plain text file."""
        path = Path(file_path)
        content = path.read_text(encoding="utf-8", errors="ignore")
        return self._parse_text_content(content, path.stem)

    def _parse_text_content(self, content: str, default_title: str = "Untitled") -> ParsedDocument:
        """Parse plain text content."""
        # Use first non-empty line as title
        lines = content.split("\n")
        title = default_title
        for line in lines:
            if line.strip():
                if len(line.strip()) < 100:
                    title = line.strip()
                break

        return ParsedDocument(
            title=title,
            content=content,
            metadata={"format": "text", "original_length": len(content)},
            links=[],
        )

    async def _parse_pdf(self, file_path: str) -> ParsedDocument:
        """Parse PDF file."""
        try:
            from pypdf import Reader
        except ImportError:
            return ParsedDocument(
                title=Path(file_path).stem,
                content="[PDF parsing not available - install pypdf]",
                metadata={"format": "pdf", "error": "pypdf not installed"},
                links=[],
            )

        path = Path(file_path)
        try:
            with open(file_path, "rb") as f:
                reader = Reader(f)
                pages = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text)

                content = "\n\n".join(pages)

                return ParsedDocument(
                    title=path.stem,
                    content=content,
                    metadata={
                        "format": "pdf",
                        "pages": len(pages),
                        "original_length": sum(len(p) for p in pages),
                    },
                    links=[],
                )
        except Exception as e:
            return ParsedDocument(
                title=path.stem,
                content=f"[PDF parsing error: {str(e)}]",
                metadata={"format": "pdf", "error": str(e)},
                links=[],
            )

    async def _parse_docx(self, file_path: str) -> ParsedDocument:
        """Parse DOCX file."""
        try:
            from docx import Document
        except ImportError:
            return ParsedDocument(
                title=Path(file_path).stem,
                content="[DOCX parsing not available - install python-docx]",
                metadata={"format": "docx", "error": "python-docx not installed"},
                links=[],
            )

        path = Path(file_path)
        try:
            doc = Document(file_path)
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Extract tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        tables_text.append(" | ".join(cells))

            content = "\n\n".join(paragraphs)
            if tables_text:
                content += "\n\n[Tables]\n" + "\n".join(tables_text)

            return ParsedDocument(
                title=path.stem,
                content=content,
                metadata={
                    "format": "docx",
                    "paragraphs": len(paragraphs),
                    "tables": len(doc.tables),
                },
                links=[],
            )
        except Exception as e:
            return ParsedDocument(
                title=path.stem,
                content=f"[DOCX parsing error: {str(e)}]",
                metadata={"format": "docx", "error": str(e)},
                links=[],
            )

    def _parse_html_content(self, content: str) -> ParsedDocument:
        """Parse HTML content."""
        # Simple HTML to text
        import html

        # Decode HTML entities
        content = html.unescape(content)

        # Remove scripts and styles
        content = re.sub(r"<script[^>]*>[\s\S]*?</script>", "", content, flags=re.IGNORECASE)
        content = re.sub(r"<style[^>]*>[\s\S]*?</style>", "", content, flags=re.IGNORECASE)

        # Extract title
        title_match = re.search(r"<title>([^<]+)</title>", content, re.IGNORECASE)
        title = title_match.group(1) if title_match else "Untitled"

        # Extract links
        links = re.findall(r'href=["\']([^"\']+)["\']', content, re.IGNORECASE)

        # Remove HTML tags
        cleaned = re.sub(r"<[^>]+>", " ", content)
        cleaned = re.sub(r"\s{2,}", " ", cleaned)

        return ParsedDocument(
            title=title,
            content=cleaned.strip(),
            metadata={"format": "html", "links_count": len(links)},
            links=links,
        )


# Singleton
_document_parser: Optional[DocumentParser] = None


def document_parser() -> DocumentParser:
    """Get document parser singleton."""
    global _document_parser
    if _document_parser is None:
        _document_parser = DocumentParser()
    return _document_parser
